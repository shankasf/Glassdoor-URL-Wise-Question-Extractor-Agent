from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
import sys
from datetime import datetime

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

class DOCXGenerator:
    def __init__(self, output_dir='scraped_data'):
        self.output_dir = output_dir
        self.document = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Setup document styles"""
        # Title style
        title_style = self.document.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Heading style
        heading_style = self.document.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
        
        # Question style
        question_style = self.document.styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
        question_style.font.name = 'Arial'
        question_style.font.size = Pt(12)
        question_style.font.bold = True
        question_style.paragraph_format.space_before = Pt(8)
        question_style.paragraph_format.space_after = Pt(4)
        
        # Answer style
        answer_style = self.document.styles.add_style('Answer', WD_STYLE_TYPE.PARAGRAPH)
        answer_style.font.name = 'Arial'
        answer_style.font.size = Pt(11)
        answer_style.paragraph_format.space_after = Pt(6)
        answer_style.paragraph_format.left_indent = Inches(0.25)
        
        # Metadata style
        metadata_style = self.document.styles.add_style('Metadata', WD_STYLE_TYPE.PARAGRAPH)
        metadata_style.font.name = 'Arial'
        metadata_style.font.size = Pt(9)
        metadata_style.font.italic = True
        metadata_style.paragraph_format.space_after = Pt(2)
    
    def create_document(self, qa_pairs, company='Unknown', position='Unknown'):
        """Create a formatted DOCX document from Q&A pairs"""
        # Add title
        title = self.document.add_paragraph(f'{company} {position} Interview Questions & Answers', style='CustomTitle')
        
        # Add metadata
        metadata = self.document.add_paragraph(
            f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n'
            f'Total Questions: {len(qa_pairs)}\n'
            f'Source: Glassdoor Interview Reviews',
            style='Metadata'
        )
        
        # Add table of contents
        self.add_table_of_contents(qa_pairs)
        
        # Add Q&A sections
        self.add_qa_sections(qa_pairs)
        
        # Add summary
        self.add_summary(qa_pairs)
        
        return self.document
    
    def add_table_of_contents(self, qa_pairs):
        """Add table of contents"""
        toc_heading = self.document.add_paragraph('Table of Contents', style='CustomHeading')
        
        for i, qa in enumerate(qa_pairs[:20]):  # Limit to first 20 for TOC
            question_preview = qa['question'][:60] + '...' if len(qa['question']) > 60 else qa['question']
            toc_item = self.document.add_paragraph(f'{i+1}. {question_preview}', style='Metadata')
    
    def add_qa_sections(self, qa_pairs):
        """Add Q&A sections to document"""
        section_heading = self.document.add_paragraph('Interview Questions & Answers', style='CustomHeading')
        
        for i, qa in enumerate(qa_pairs):
            # Question
            question_para = self.document.add_paragraph(f'Q{i+1}: {qa["question"]}', style='Question')
            
            # Answer
            answer_para = self.document.add_paragraph(f'Answer: {qa["answer"]}', style='Answer')
            
            # Metadata
            metadata_items = []
            if qa.get('difficulty'):
                metadata_items.append(f'Difficulty: {qa["difficulty"]}')
            if qa.get('outcome'):
                metadata_items.append(f'Outcome: {qa["outcome"]}')
            if qa.get('location'):
                metadata_items.append(f'Location: {qa["location"]}')
            if qa.get('date'):
                metadata_items.append(f'Date: {qa["date"]}')
            
            if metadata_items:
                metadata_para = self.document.add_paragraph(' | '.join(metadata_items), style='Metadata')
            
            # Add separator
            if i < len(qa_pairs) - 1:
                self.document.add_paragraph('─' * 80, style='Metadata')
    
    def add_summary(self, qa_pairs):
        """Add summary section"""
        summary_heading = self.document.add_paragraph('Summary', style='CustomHeading')
        
        # Calculate statistics
        companies = set(qa.get('company', 'Unknown') for qa in qa_pairs)
        positions = set(qa.get('position', 'Unknown') for qa in qa_pairs)
        difficulties = [qa.get('difficulty', '') for qa in qa_pairs if qa.get('difficulty')]
        outcomes = [qa.get('outcome', '') for qa in qa_pairs if qa.get('outcome')]
        
        # Difficulty distribution
        difficulty_dist = {}
        for diff in difficulties:
            difficulty_dist[diff] = difficulty_dist.get(diff, 0) + 1
        
        # Outcome distribution
        outcome_dist = {}
        for outcome in outcomes:
            outcome_dist[outcome] = outcome_dist.get(outcome, 0) + 1
        
        # Add statistics
        stats_para = self.document.add_paragraph(
            f'Total Questions: {len(qa_pairs)}\n'
            f'Companies: {", ".join(companies)}\n'
            f'Positions: {", ".join(positions)}\n'
            f'Difficulty Distribution: {dict(difficulty_dist)}\n'
            f'Outcome Distribution: {dict(outcome_dist)}',
            style='Answer'
        )
    
    def save_document(self, filename=None):
        """Save the document to file"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f'interview_questions_{timestamp}.docx'
        
        filepath = os.path.join(self.output_dir, filename)
        self.document.save(filepath)
        return filepath

def generate_docx_from_qa(qa_pairs, company='Unknown', position='Unknown', output_dir='scraped_data'):
    """Generate DOCX file from Q&A pairs"""
    generator = DOCXGenerator(output_dir)
    generator.create_document(qa_pairs, company, position)
    filepath = generator.save_document()
    return filepath

def main():
    # This would be called from the main scraper
    pass

if __name__ == "__main__":
    main()
